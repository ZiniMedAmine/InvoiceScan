import json
from docx import Document
import csv


def parse_text_to_csv(text_file):
  """
  Parses text data from a file into a dictionary structure, handling potential JSON parsing errors.

  Args:
      text_file (str): Path to the text file containing data.

  Returns:
      tuple: A tuple containing either (None, None) if parsing fails 
             or (document_type, extracted_data) if parsing is successful.
  """
  try:
      with open(text_file, 'r', encoding='utf-8') as f:
          data = json.load(f)
      document_type = "Auto"

      extracted_data = {}
      def flatten_data(value, parent_key=""):
          """
          Flattens nested dictionaries recursively, handling all levels and 
          resulting in single-level keys and values.

          Args:
              value: The value to flatten (can be a dictionary or basic type).
              parent_key (str, optional): The parent key for nested elements (defaults to "").
          """
          if isinstance(value, dict):
              for key, item in value.items():
                  new_key = parent_key + "." + key if parent_key else key
                  flatten_data(item, new_key)
          else:
              # Handle basic data types (strings, numbers, booleans)
              extracted_data[parent_key] = value

      flatten_data(data)
      return document_type, extracted_data
  except json.decoder.JSONDecodeError:
      print("Error: Invalid JSON data format in file", text_file)
      return None, None


def write_to_csv(data, csv_file):
  """
  Writes parsed data to a CSV file, handling flattened data structure.

  Args:
      data (tuple): A tuple containing the document type (str) and extracted data (dict).
      csv_file (str): Path to the output CSV file.
  """
  if data is None:
      # Handle parsing failure (create empty CSV)
      with open(csv_file, 'w', newline='', encoding="utf-8") as f:
          writer = csv.writer(f)
          writer.writerow(["Error: Failed to parse data"])
      return

  document_type, extracted_data = data

  # Write document type as the first row
  with open(csv_file, 'w', newline='', encoding="utf-8") as f:
      writer = csv.writer(f)
      writer.writerow([f"Document Type: {document_type}"])

  # Write flattened data as separate rows
  with open(csv_file, 'a', newline='', encoding="utf-8") as f:
      writer = csv.writer(f)
      for key, value in extracted_data.items():
          writer.writerow([key, value])

def json_to_text(data, document_path):
    """
    Converts JSON data into a formatted Word document.

    Args:
        data (dict): The JSON data to be converted.
        document_path (str): The path to save the Word document.

    Returns:
        str: Path to the saved Word document.
    """
    # Create a new Word document
    doc = Document()
    
    # Loop through the JSON data and add content to the document
    for section_title, section_content in data.items():
        # Add section title as a heading
        doc.add_heading(section_title, level=1)
        
        # Handle nested dictionaries
        if isinstance(section_content, dict):
            for key, value in section_content.items():
                # Add key-value pairs as paragraphs
                doc.add_paragraph(f"{key}: {value}")
        elif isinstance(section_content, list):
            # Handle lists
            for item in section_content:
                doc.add_paragraph(item)
        else:
            doc.add_paragraph(str(section_content))
    
    # Save the document to the specified path
    doc.save(document_path)
    
    return document_path